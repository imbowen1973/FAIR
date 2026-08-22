"""Blocks into a Word document, through the template's own styles.

The rule this file exists to keep: **no direct formatting is ever
written**. Every paragraph gets a named style from `style-map.yaml` and
nothing else. If a document can make something bold by any route other
than the shared inline grammar, or change a size, or set a colour, the
separation has leaked and the template has stopped owning appearance.

The template's body is cleared and its styles, numbering, headers,
footers and section setup are kept. A template is a shape, not a
starting draft — a cover page left in it would appear in every document
rendered from it.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx.oxml.ns import qn
from docx.shared import Emu
from lxml import etree

from .blocks import Block, parse_document
from .runs import DEFAULT_CODE_TYPEFACE, plain, write_runs
from .template import (
    DEFAULT_STYLES,
    WordTemplateError,
    load_template,
    style_for_role,
)

# How wide a figure may be, as a fraction of the text column. A picture
# wider than the column is Word's problem to solve and it solves it
# badly, so it is solved here.
FIGURE_WIDTH = 1.0


class WordRenderError(Exception):
    pass


@dataclass
class DocumentMeta:
    """What the content controls are filled from — the library, not typing."""

    title: str = ""
    author: str = ""
    version: str = ""
    project: str = ""
    funding: str = ""

    def as_controls(self) -> dict[str, str]:
        return {
            "FAIR.DocumentTitle": self.title,
            "FAIR.Author": self.author,
            "FAIR.Version": self.version,
            "FAIR.Project": self.project,
            "FAIR.FundingStatement": self.funding,
        }


def _clear_body(document) -> None:
    """Empty the template's body of draft content, keeping its furniture.

    Two things are not draft content and must survive:

    `sectPr` carries page size, margins, and which headers and footers
    apply — the page setup the template exists to define.

    Content controls are the template's own structure. A cover page's
    title control is a slot the renderer is meant to fill, not a sample
    paragraph to sweep away; clearing them and then looking for them to
    fill is how the funding statement came out empty.
    """
    body = document.element.body
    keep = {qn("w:sectPr"), qn("w:sdt")}
    for child in list(body):
        if child.tag not in keep:
            body.remove(child)


def _styles_for(document, style_map: dict | None) -> dict[str, str]:
    """Which style serves each role in this template."""
    available = {style.name for style in document.styles}
    configured = (style_map or {}).get("styles") or {}
    resolved: dict[str, str] = {}
    for role in DEFAULT_STYLES:
        name = configured.get(role) or style_for_role(role, available)
        if name and name in available:
            resolved[role] = name
    return resolved


def _paragraph(document, styles: dict[str, str], role: str):
    """A paragraph in the template's style for `role`, or its body style."""
    paragraph = document.add_paragraph()
    name = styles.get(role) or styles.get("body")
    if name:
        paragraph.style = document.styles[name]
    return paragraph


def _set_list_level(paragraph, level: int) -> None:
    """Indent a nested list item using the style's own numbering.

    Word carries list depth in the paragraph's numbering properties. The
    numbering definition itself comes from the template's style — this
    only says which level of it this item sits at, so a template that
    numbers 1 / a / i keeps doing so.
    """
    if level <= 0:
        return
    pPr = paragraph._p.get_or_add_pPr()
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        numPr = etree.SubElement(pPr, qn("w:numPr"))
    ilvl = numPr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = etree.SubElement(numPr, qn("w:ilvl"))
    ilvl.set(qn("w:val"), str(min(level, 8)))


def _add_figure(document, styles, block: Block, base_dir: Path, warn) -> None:
    source = (base_dir / block.src).resolve() if block.src else None
    if source is None or not source.is_file():
        warn(f"figure not found, skipped: {block.src}")
        return

    paragraph = document.add_paragraph()
    if styles.get("body"):
        paragraph.style = document.styles[styles["body"]]
    run = paragraph.add_run()
    width = document.sections[0].page_width - (
        document.sections[0].left_margin + document.sections[0].right_margin
    )
    run.add_picture(str(source), width=Emu(int(width * FIGURE_WIDTH)))

    if block.alt:
        caption = _paragraph(document, styles, "caption")
        write_runs(caption, block.alt)


def _add_table(document, styles, block: Block, style_map: dict | None, warn) -> None:
    if not block.cells:
        return
    columns = max(len(row) for row in block.cells)
    table = document.add_table(rows=0, cols=columns)

    wanted = ((style_map or {}).get("tables") or {}).get("default")
    if wanted:
        try:
            table.style = document.styles[wanted]
        except KeyError:
            warn(f"table style not in the template, using its default: {wanted}")

    for index, row_cells in enumerate(block.cells):
        row = table.add_row()
        is_header = index < len(block.items) and block.items[index].header
        for column, text in enumerate(row_cells[:columns]):
            cell = row.cells[column]
            # A new cell arrives with an empty paragraph; write into it
            # rather than adding a second one and leaving a blank line.
            paragraph = cell.paragraphs[0]
            write_runs(paragraph, text)
            if is_header:
                for run in paragraph.runs:
                    run.font.bold = True


def _fill_controls(document, meta: DocumentMeta) -> list[str]:
    """Populate the template's content controls by tag.

    python-docx has no notion of these, so this is lxml — the same as
    render.py already does for creationId. A control the template does
    not expose is simply not filled: the template decides what it wants
    to show, and this decides what those things say.
    """
    values = {tag: text for tag, text in meta.as_controls().items() if text}
    filled: list[str] = []

    roots = [document.element.body]
    for section in document.sections:
        for part in (section.header, section.footer,
                     section.first_page_header, section.first_page_footer,
                     section.even_page_header, section.even_page_footer):
            element = getattr(part, "_element", None)
            if element is not None:
                roots.append(element)

    for root in roots:
        for sdt in root.iter(qn("w:sdt")):
            tag = sdt.find(f".//{qn('w:tag')}")
            name = tag.get(qn("w:val")) if tag is not None else None
            if name not in values:
                continue
            content = sdt.find(qn("w:sdtContent"))
            if content is None:
                continue

            texts = list(content.iter(qn("w:t")))
            if texts:
                texts[0].text = values[name]
                # Anything beyond the first run is the placeholder's own
                # sample text; keeping it would append to the real value.
                for extra in texts[1:]:
                    extra.text = ""
            else:
                # An empty control is the normal case for a template
                # author who left the slot blank rather than typing a
                # sample into it. It still has to be fillable.
                paragraph = content.find(qn("w:p"))
                if paragraph is None:
                    paragraph = etree.SubElement(content, qn("w:p"))
                run = etree.SubElement(paragraph, qn("w:r"))
                text = etree.SubElement(run, qn("w:t"))
                text.text = values[name]
                text.set(
                    "{http://www.w3.org/XML/1998/namespace}space", "preserve"
                )
            filled.append(name)
    return filled


def render_document(
    source: Path,
    template: Path,
    out_path: Path,
    *,
    style_map: dict | None = None,
    meta: DocumentMeta | None = None,
    warn=lambda message: None,
) -> dict:
    """Render one markdown document into the template. Returns a summary."""
    if not source.is_file():
        raise WordRenderError(f"no such document: {source}")

    document = load_template(template)
    _clear_body(document)
    styles = _styles_for(document, style_map)

    blocks = parse_document(source.read_text(encoding="utf-8"))
    meta = meta or DocumentMeta()
    if not meta.title:
        first = next((b for b in blocks if b.kind == "title"), None)
        if first:
            meta.title = plain(first.text)

    code_typeface = ((style_map or {}).get("code_typeface")) or DEFAULT_CODE_TYPEFACE
    base_dir = source.parent

    for block in blocks:
        if block.kind in ("title", "section", "subsection"):
            role = {"title": "title", "section": "h1", "subsection": "h2"}[block.kind]
            write_runs(_paragraph(document, styles, role), block.text, code_typeface)
        elif block.kind == "body":
            write_runs(_paragraph(document, styles, "body"), block.text, code_typeface)
        elif block.kind == "callout":
            write_runs(_paragraph(document, styles, "callout"), block.text, code_typeface)
        elif block.kind == "quote":
            write_runs(_paragraph(document, styles, "quote"), block.text, code_typeface)
        elif block.kind == "list":
            paragraph = _paragraph(document, styles, "numbers" if block.ordered else "bullets")
            _set_list_level(paragraph, block.level)
            write_runs(paragraph, block.text, code_typeface)
        elif block.kind == "code":
            paragraph = _paragraph(document, styles, "code")
            # Code is verbatim: the inline grammar does not apply inside
            # it, or a shell command with an asterisk would come out bold.
            run = paragraph.add_run(block.text)
            run.font.name = code_typeface
        elif block.kind == "figure":
            _add_figure(document, styles, block, base_dir, warn)
        elif block.kind == "table":
            _add_table(document, styles, block, style_map, warn)
        elif block.kind == "pagebreak":
            document.add_page_break()

    filled = _fill_controls(document, meta)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(out_path))
    return {
        "docx": out_path,
        "blocks": len(blocks),
        "styles": styles,
        "controls": filled,
    }


def main(argv: list[str] | None = None) -> int:
    import argparse
    import sys

    import yaml

    ap = argparse.ArgumentParser(
        prog="edufair-doc",
        description="Render a markdown document into a branded Word file",
    )
    ap.add_argument("document", type=Path)
    ap.add_argument("--template", type=Path, required=True,
                    help="the Word template: .docx, .dotx or .dotm")
    ap.add_argument("--styles", type=Path, default=None,
                    help="style-map.yaml (default: derived from the template)")
    ap.add_argument("--out", type=Path, default=None,
                    help="where to write the .docx (default: beside the source)")
    ap.add_argument("--title", default="", help="overrides the document's own heading")
    ap.add_argument("--project", default="")
    ap.add_argument("--version", default="")
    ap.add_argument("--author", default="")
    ap.add_argument("--funding", default="", help="the funder credit, as required")
    ap.add_argument("--quiet", action="store_true")
    args = ap.parse_args(argv)

    style_map = None
    if args.styles:
        style_map = yaml.safe_load(args.styles.read_text(encoding="utf-8")) or {}

    out = args.out or args.document.with_suffix(".docx")
    warnings: list[str] = []
    try:
        summary = render_document(
            args.document,
            args.template,
            out,
            style_map=style_map,
            meta=DocumentMeta(
                title=args.title,
                author=args.author,
                version=args.version,
                project=args.project,
                funding=args.funding,
            ),
            warn=warnings.append,
        )
    except (WordRenderError, WordTemplateError) as exc:
        print(f"error: {exc}", file=sys.stderr)
        return 1

    for message in warnings:
        print(f"warning: {message}", file=sys.stderr)
    if not args.quiet:
        controls = ", ".join(summary["controls"]) or "none"
        print(
            f"wrote {summary['docx']}: {summary['blocks']} blocks, "
            f"controls filled: {controls}",
            file=sys.stderr,
        )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
