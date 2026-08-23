"""The Word renderer: loading a template and reading what it offers."""

from __future__ import annotations

import re
import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

from edufair_renderer.word.render import render_document
from edufair_renderer.word.template import (
    DOCUMENT_CT,
    TEMPLATE_CTS,
    WordTemplateError,
    control_tags,
    draft_style_map,
    inspect_template,
    load_template,
    style_for_role,
)


def _branded(path: Path, *, styles=(), controls=()) -> Path:
    """A Word file with named styles and content controls, as a project would brand one."""
    document = Document()
    for name in styles:
        document.styles.add_style(name, 1)  # 1 = paragraph style
    body = document.element.body
    for tag_value in controls:
        sdt = etree.SubElement(body, qn("w:sdt"))
        properties = etree.SubElement(sdt, qn("w:sdtPr"))
        tag = etree.SubElement(properties, qn("w:tag"))
        tag.set(qn("w:val"), tag_value)
        content = etree.SubElement(sdt, qn("w:sdtContent"))
        etree.SubElement(content, qn("w:p"))
    document.save(path)
    return path


def _retype(src: Path, dst: Path, content_type: str) -> Path:
    """The same package, typed as a template rather than a document."""
    with zipfile.ZipFile(src) as zin, zipfile.ZipFile(dst, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "[Content_Types].xml":
                data = data.replace(DOCUMENT_CT.encode(), content_type.encode())
            zout.writestr(item, data)
    return dst


@pytest.mark.parametrize("content_type", TEMPLATE_CTS[:2], ids=["dotx", "dotm"])
def test_a_template_opens_whatever_word_saved_it_as(tmp_path, content_type):
    """python-docx refuses a template outright; the loader retypes it first.

    Which flavour the branding arrives in is not something an author
    should have to think about.
    """
    src = _branded(tmp_path / "brand.docx", styles=["FAIR Callout"])
    template = _retype(src, tmp_path / "brand.template", content_type)

    document = load_template(template)
    assert "FAIR Callout" in [s.name for s in document.styles], (
        "the template's own styles must survive the retype — they are the point of it"
    )


def test_a_plain_docx_opens_too(tmp_path):
    src = _branded(tmp_path / "brand.docx", styles=["FAIR Callout"])
    assert "FAIR Callout" in [s.name for s in load_template(src).styles]


def test_an_old_binary_doc_is_refused_with_a_way_out(tmp_path):
    old = tmp_path / "handbook.doc"
    old.write_bytes(b"\xd0\xcf\x11\xe0 an OLE2 document from 2003")
    with pytest.raises(WordTemplateError, match="save as .docx"):
        load_template(old)


def test_content_controls_are_found_by_tag(tmp_path):
    # The tag is the contract: the renderer fills FAIR.DocumentTitle
    # wherever the template chose to put it.
    src = _branded(
        tmp_path / "brand.docx",
        controls=["FAIR.DocumentTitle", "FAIR.FundingStatement"],
    )
    assert control_tags(load_template(src)) == [
        "FAIR.DocumentTitle",
        "FAIR.FundingStatement",
    ]


def test_a_template_brands_a_role_by_naming_a_style_for_it(tmp_path):
    """Define "FAIR Callout" in Word and callouts use it — no map to edit."""
    src = _branded(tmp_path / "brand.docx", styles=["FAIR Callout", "FAIR Body"])
    styles = draft_style_map(inspect_template(src))["styles"]

    assert styles["callout"] == "FAIR Callout"
    assert styles["body"] == "FAIR Body"
    # Roles the project has not branded fall back to Word's own.
    assert styles["h1"] == "Heading 1"
    assert styles["quote"] == "Quote"


def test_a_role_the_template_cannot_serve_is_left_out_not_guessed(tmp_path):
    # A map naming a style that is not there renders wrong rather than
    # failing: Word silently falls back to Normal.
    available = {"Heading 1", "Normal"}
    assert style_for_role("h1", available) == "Heading 1"
    assert style_for_role("callout", available) is None

    src = _branded(tmp_path / "brand.docx")
    report = inspect_template(src)
    styles = draft_style_map(report)["styles"]
    for role, name in styles.items():
        assert name in report.paragraph_styles, f"{role} points at a style that is not there"


def test_an_explicit_map_is_checked_exactly_as_written(tmp_path):
    """A map is a promise, so it is checked against the template as given."""
    src = _branded(tmp_path / "brand.docx")
    report = inspect_template(src, {"styles": {"callout": "House Callout"}})
    assert report.missing.get("callout") == "House Callout"
    assert not report.ok


def test_a_plain_word_template_still_renders(tmp_path):
    """A project that has branded nothing must still get a document."""
    src = _branded(tmp_path / "plain.docx")
    styles = draft_style_map(inspect_template(src))["styles"]
    for role in ("title", "h1", "h2", "body", "quote", "caption", "bullets", "numbers"):
        assert role in styles, f"Word's own default for {role} should have been found"


# ---- markdown to blocks -------------------------------------------------


def _blocks(text):
    from edufair_renderer.word.blocks import parse_document

    return parse_document(text)


def test_a_document_becomes_the_blocks_it_is_made_of():
    blocks = _blocks(
        "# Title\n\nBody.\n\n## Section\n\n- one\n  - nested\n\n"
        "1. first\n\n> quoted\n\n```py\nx=1\n```\n\n---\n"
    )
    assert [b.kind for b in blocks] == [
        "title", "body", "section",
        "list", "list", "list",
        "quote", "code", "pagebreak",
    ]
    nested = blocks[4]
    assert nested.level == 1, "a nested item keeps its depth"
    assert blocks[5].ordered is True and blocks[3].ordered is False


def test_an_image_alone_is_a_figure_and_one_in_a_sentence_is_not():
    figure = _blocks("![The cold chain](media/chain.png)\n")[0]
    assert figure.kind == "figure"
    assert figure.src == "media/chain.png" and figure.alt == "The cold chain"

    # Inside a sentence it is inline, and stays in the text for the run
    # writer to deal with.
    inline = _blocks("See ![this](x.png) here.\n")[0]
    assert inline.kind == "body"


def test_a_table_keeps_which_row_is_the_header():
    table = _blocks("| Time | Activity |\n|---|---|\n| 0-10 | Opening |\n")[0]
    assert table.kind == "table"
    assert table.cells == [["Time", "Activity"], ["0-10", "Opening"]]
    assert [item.header for item in table.items] == [True, False]


def test_a_generated_outcomes_section_becomes_a_callout():
    # The workbench writes these markers around the outcomes it generates,
    # and the section is the session's own claim -- it reads as one.
    blocks = _blocks(
        "<!-- outcomes: from outcomes.yaml -->\n\n"
        "By the end of the session:\n\n"
        "<!-- /outcomes -->\n\nOrdinary text.\n"
    )
    assert blocks[0].kind == "callout"
    assert blocks[1].kind == "body"


def test_the_block_model_has_nowhere_to_put_formatting():
    """The separation, enforced by absence rather than by a rule."""
    from edufair_renderer.word.blocks import Block

    assert not any(
        field in Block.__dataclass_fields__
        for field in ("font", "colour", "color", "size", "margin", "align")
    )


# ---- rendering ----------------------------------------------------------


def _render(tmp_path, markdown, *, template=None, meta=None, files=()):
    from edufair_renderer.word.render import DocumentMeta, render_document

    tpl = template or _branded(tmp_path / "brand.docx")
    source = tmp_path / "doc.md"
    source.write_text(markdown, encoding="utf-8")
    for name, maker in files:
        maker(tmp_path / name)
    out = tmp_path / "out.docx"
    warnings = []
    render_document(
        source, tpl, out,
        style_map=draft_style_map(inspect_template(tpl)),
        meta=meta or DocumentMeta(),
        warn=warnings.append,
    )
    return Document(out), warnings, out


def test_each_block_gets_the_template_style_for_its_role(tmp_path):
    document, _, _ = _render(
        tmp_path,
        "# Handbook\n\nBody.\n\n## Section\n\n- one\n\n1. first\n\n> quoted\n",
    )
    styled = [(p.style.name, p.text) for p in document.paragraphs if p.text.strip()]
    assert styled == [
        ("Title", "Handbook"),
        ("Normal", "Body."),
        ("Heading 1", "Section"),
        ("List Bullet", "one"),
        ("List Number", "first"),
        ("Quote", "quoted"),
    ]


def test_the_same_markdown_through_two_templates_says_the_same_thing(tmp_path):
    """The separation, which is the whole point.

    Same words in the same order; a different-looking file. If markdown
    could change a font, this would not hold.
    """
    from docx.shared import Pt, RGBColor

    def branded(path, size, colour):
        document = Document()
        document.styles["Normal"].font.size = Pt(size)
        document.styles["Normal"].font.color.rgb = RGBColor.from_string(colour)
        document.save(path)
        return path

    markdown = "# Handbook\n\nBody text.\n\n## Section\n\n- one\n"
    words = []
    for name, size, colour in (("a", 11, "222222"), ("b", 16, "AA0000")):
        sub = tmp_path / name
        sub.mkdir()
        document, _, _ = _render(
            sub, markdown, template=branded(sub / "brand.docx", size, colour)
        )
        words.append([p.text for p in document.paragraphs if p.text.strip()])
        assert document.styles["Normal"].font.size.pt == size

    assert words[0] == words[1]


def test_no_direct_formatting_reaches_the_document(tmp_path):
    """Appearance belongs to the template, and nothing here may write it."""
    _, _, out = _render(tmp_path, "# T\n\nBody with **bold**.\n\n## S\n\n- one\n")
    with zipfile.ZipFile(out) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    for leaked in ("<w:sz ", "<w:color ", "<w:jc ", "<w:spacing ", "<w:ind "):
        assert leaked not in xml, f"{leaked!r} is the template's business, not ours"


def test_inline_marks_are_written_as_word_run_properties(tmp_path):
    _, _, out = _render(tmp_path, "H~2~O at x^2^, ~~struck~~ and a [link](https://e.org).\n")
    with zipfile.ZipFile(out) as archive:
        xml = archive.read("word/document.xml").decode("utf-8")
    assert 'w:val="subscript"' in xml
    assert 'w:val="superscript"' in xml
    assert "<w:strike" in xml
    # A real hyperlink, not blue text that does nothing when clicked.
    assert "w:hyperlink" in xml


def test_the_templates_cover_page_does_not_appear_in_every_document(tmp_path):
    document = Document()
    document.add_paragraph("SAMPLE COVER — a draft, not a shape")
    tpl = tmp_path / "brand.docx"
    document.save(tpl)

    rendered, _, _ = _render(tmp_path, "# Real\n\nContent.\n", template=tpl)
    assert "SAMPLE COVER" not in "\n".join(p.text for p in rendered.paragraphs)


def test_content_controls_survive_the_clear_and_are_filled(tmp_path):
    """They are the template's structure, not its draft content.

    Clearing them and then looking for them to fill is how the funding
    statement came out empty.
    """
    from edufair_renderer.word.render import DocumentMeta

    tpl = _branded(
        tmp_path / "brand.docx",
        controls=["FAIR.DocumentTitle", "FAIR.FundingStatement"],
    )
    document, _, _ = _render(
        tmp_path,
        "# The handbook\n\nBody.\n",
        template=tpl,
        meta=DocumentMeta(funding="Funded by the European Union"),
    )
    xml = document.element.body.xml
    assert "Funded by the European Union" in xml
    # The title comes from the document's own heading, not from typing.
    assert "The handbook" in xml


def test_a_figure_is_embedded_and_a_missing_one_is_reported(tmp_path):
    from PIL import Image

    def make(path):
        Image.new("RGB", (800, 600), (200, 60, 60)).save(path)

    document, warnings, out = _render(
        tmp_path,
        "![The cold chain](chain.png)\n\n![gone](nowhere.png)\n",
        files=[("chain.png", make)],
    )
    with zipfile.ZipFile(out) as archive:
        embedded = [n for n in archive.namelist() if n.startswith("word/media/")]
    assert len(embedded) == 1
    assert any("nowhere.png" in w for w in warnings), "a missing figure must be said, not silent"
    # The alt text becomes the caption, in the template's caption style.
    assert any(p.style.name == "Caption" and "cold chain" in p.text for p in document.paragraphs)


def test_a_table_arrives_as_a_table(tmp_path):
    document, _, _ = _render(
        tmp_path, "| Time | Activity |\n|---|---|\n| 0-10 | Opening |\n"
    )
    assert len(document.tables) == 1
    table = document.tables[0]
    assert [cell.text for cell in table.rows[0].cells] == ["Time", "Activity"]
    assert [cell.text for cell in table.rows[1].cells] == ["0-10", "Opening"]


def test_code_is_verbatim(tmp_path):
    # The inline grammar does not apply inside a code block, or a shell
    # command with an asterisk would come out bold.
    document, _, _ = _render(tmp_path, "```sh\nrm *.tmp && echo **done**\n```\n")
    assert any("rm *.tmp && echo **done**" in p.text for p in document.paragraphs)


# ---- instructions written from a library --------------------------------


def test_the_prompt_carries_this_template_and_no_other(tmp_path):
    """A model told only about "the grammar" invents a TwoColumn layout.

    The prompt is generated so it names what this template actually has.
    """
    import shutil

    from edufair_renderer.prompt import build_prompt

    root = tmp_path / "lib"
    (root / "competencies").mkdir(parents=True)
    shutil.copy(
        Path(__file__).resolve().parents[2] / "examples" / "layout-map.yaml",
        root / "layout-map.yaml",
    )
    (root / "competencies" / "framework.yaml").write_text(
        "competencies:\n  CE1:\n    label: Designing clinical teaching\n",
        encoding="utf-8",
    )
    (root / "course.yaml").write_text("id: c\ntitle: A course\n", encoding="utf-8")

    prompt = build_prompt(root)

    # Every layout the template has, with its own regions.
    assert "`Comparison`" in prompt
    assert "`left_head`" in prompt
    # The course's own competency ids, so `develops:` comes back valid.
    assert "`CE1` — Designing clinical teaching" in prompt


def test_the_prompt_makes_the_model_ask_before_it_writes():
    """A topic alone tells a writer nothing about what the room needs.

    Subject, audience, duration and shape are the prompt, not a
    preamble to it.
    """
    from edufair_renderer.prompt import generic_prompt

    prompt = generic_prompt()
    assert prompt, "the generic half must be findable"
    assert "Ask before you write" in prompt
    for question in ("The subject", "The audience", "How long", "The structure"):
        assert question in prompt, f"the intake must ask about {question.lower()}"
    # And the rules that matter most, stated as rules.
    assert "never describe how anything looks" in prompt.lower()
    assert "How long is how many slides" in prompt


def test_the_library_half_sits_under_the_generic_one(tmp_path):
    """Both halves, in that order: how to write, then what with."""
    import shutil

    from edufair_renderer.prompt import build_prompt

    root = tmp_path / "lib"
    root.mkdir()
    shutil.copy(
        Path(__file__).resolve().parents[2] / "examples" / "layout-map.yaml",
        root / "layout-map.yaml",
    )
    prompt = build_prompt(root)
    assert prompt.index("Ask before you write") < prompt.index("This library:")

    # generic=False is the library half on its own, for appending to a
    # prompt somebody has already edited by hand.
    alone = build_prompt(root, generic=False)
    assert "This library:" in alone
    assert "Ask before you write" not in alone


def test_the_worked_example_uses_a_layout_the_template_has(tmp_path):
    """An example naming a layout that does not exist teaches the wrong thing."""
    import re
    import shutil

    from edufair_renderer.layoutmap import load_layout_map
    from edufair_renderer.prompt import build_prompt

    root = tmp_path / "lib"
    root.mkdir()
    shutil.copy(
        Path(__file__).resolve().parents[2] / "examples" / "layout-map.yaml",
        root / "layout-map.yaml",
    )
    prompt = build_prompt(root)
    layouts = set(load_layout_map(root / "layout-map.yaml"))

    used = re.findall(r"^layout: (\w+)$", prompt, re.MULTILINE)
    assert used, "the prompt must show a worked slide"
    for name in used:
        assert name in layouts, f"the example uses {name}, which this template lacks"


# ---- rendering needs no template ---------------------------------------
#
# This is the design, not a convenience. A .docx carries its own
# styles.xml, and Word's built-in styles are defined against the theme
# rather than against literal fonts and colours -- so a document that
# names built-in styles takes its appearance from the reader's own Word.
# Requiring a template would mean shipping and keeping one in step for no
# gain, and would put a file on a server that never needed to be there.

def test_a_document_renders_with_no_template_at_all(tmp_path):
    source = tmp_path / "workbook.md"
    source.write_text(
        "# Cold chain\n\n"
        "Vaccines fail quietly.\n\n"
        "## What to check\n\n"
        "- The log, not the light\n"
        "- The door seal\n",
        encoding="utf-8",
    )
    out = tmp_path / "workbook.docx"
    summary = render_document(source, None, out)

    assert out.is_file(), "no template is the ordinary case, not an error"
    assert summary["blocks"] > 0

    document = Document(str(out))
    used = [p.style.name for p in document.paragraphs if p.text.strip()]
    assert "Title" in used
    assert any(name.startswith("Heading") for name in used), (
        "headings must be real styles or they miss Word's navigation pane"
    )


def test_the_built_in_styles_follow_the_theme_not_a_literal(tmp_path):
    """The reason no template is needed, pinned.

    If these ever became literal fonts and colours, a rendered document
    would stop taking its appearance from the reader's Word and a
    template would quietly become necessary again.
    """
    source = tmp_path / "d.md"
    source.write_text("# T\n\n## H\n\nWords.\n", encoding="utf-8")
    out = tmp_path / "d.docx"
    render_document(source, None, out)

    xml = Document(str(out)).styles.element.xml
    for style_id in ("Heading1", "Title"):
        match = re.search(
            r'<w:style [^>]*w:styleId="%s".*?</w:style>' % style_id, xml, re.S
        )
        assert match, f"{style_id} is not defined"
        frag = match.group(0)
        assert "Theme" in frag or "themeColor" in frag, (
            f"{style_id} names a literal typeface or colour, so the document "
            "no longer follows the reader's theme"
        )


def test_a_template_is_still_honoured_when_there_is_one(tmp_path):
    """Optional does not mean ignored: a house style still binds."""
    brand = _branded(tmp_path / "brand.docx", styles=["eduFAIR Callout"])
    source = tmp_path / "d.md"
    source.write_text("# T\n\nWords.\n", encoding="utf-8")

    out = tmp_path / "d.docx"
    render_document(source, brand, out)
    names = {s.name for s in Document(str(out)).styles}
    assert "eduFAIR Callout" in names, "the template's own styles must survive"
